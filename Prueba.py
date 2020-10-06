   #https://python-docx.readthedocs.io/en/latest/index.html
from docx import Document
from docx.shared import Inches

from docx.enum.section import WD_ORIENT

from docx.shared import Mm
from docx.shared import Cm

import pyodbc

#Función para quitar caracteres no admitidos
def depurar (lista):
    for i in range (len(lista)):
        while (lista[i].find("  ")!=-1):
            lista[i]=lista[i].replace("  ","")
        cadena_aux=lista[i]
        lista_aux=list(cadena_aux)
        if lista_aux[-1]==" ":
            lista_aux.pop(-1)
        cadena_aux="".join(lista_aux)
        lista[i]=cadena_aux

##PARA SACAR LOS DATOS DIRECTAMENTE DE MICROSOFT ACCESS
#https://datatofish.com/how-to-connect-python-to-ms-access-database-using-pyodbc/
#ESPECIFICAR
conn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb)};DBQ=C:\Users\javie\Documents\Programas\be.mdb;')
cursor = conn.cursor()
cursor.execute("select * from Productos")
cont=0
antibiotico=[]
sustancia=[]
existentes=[]
for row in cursor.fetchall():
    if (row[1].find("A1.")!=(-1)) or (row[1].find("A2.")!=(-1)):
        cont+=1
        antibiotico.append(row[1])
        sustancia.append(row[6])
        existentes.append(row[13])
depurar(antibiotico)
depurar(sustancia)
lista_conjunta=[]
for i in range (cont):
    lista_conjunta.append([])
    for j in range (3):
        lista_conjunta[i].append("")
cont = 0
for i in range (len(antibiotico)):
    lista_conjunta[cont][0]=antibiotico[cont]
    lista_conjunta[cont][1]=sustancia[cont]
    lista_conjunta[cont][2]=existentes[cont]
    cont+=1
lista_conjunta.sort()

lista=[]

for i in range (len(lista_conjunta)):
    cadena_aux=str(lista_conjunta[i][0])+"                          SUST. ACTIVA: "+ str(lista_conjunta[i][1])
    cadena_aux=cadena_aux.replace("['","")
    cadena_aux=cadena_aux.replace("']","")
    lista.append(cadena_aux)

#PARA CUANDO SE QUIERA SACAR DE UNA HOJA DE WORD CON EL FORMATO ESTABLECIDO
##https://www.geeksforgeeks.org/python-working-with-docx-module/
##ESPECIFICAR
#doc = Document("E:\SAN XAVIER\INFORMACION\Archivos Salubridad\LISTADO DE ANTIBIOTICOS GENERICO (A2) HEROICO.docx")
#tables = doc.tables
#cont = 0
#lista = []
#for table in tables:
#    for row in table.rows:
#        for cell in row.cells:
#            for paragraph in cell.paragraphs:
#                lista.append(paragraph.text)


document = Document()

cont=0
for i in lista:

    section = document.sections[-1]

    #https://es.stackoverflow.com/questions/318813/c%c3%b3mo-puedo-cambiar-la-orientaci%c3%b3n-de-un-documento-word-usando-python-docx/318815#318815
    #https://stackoverflow.com/questions/43724030/how-to-change-page-size-to-a4-in-python-docx
    section.page_height = Mm(355.6)
    section.page_width = Mm(215.9)
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    #https://stackoverflow.com/questions/32914595/modify-docx-page-margins-with-python-docx?newreg=42213d72873f48d6b9f1f6bda4b6e4a0
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    #https://www.geeksforgeeks.org/python-working-with-docx-module/
    #https://stackoverflow.com/questions/24031011/python-docx-library-text-align
    document.add_paragraph("No.______").alignment = 2
    
    linea_arriba = document.add_table(1,2)
    linea_arriba.columns[0].cells[0].width = Cm(29.19)
    linea_arriba.columns[1].cells[0].width = Cm(4.75)
    linea_arriba.rows[0].cells[0].text = 'ESTA PÁGINA ESTÁ DEDICADA A: ' + i
    linea_arriba.rows[0].cells[1].text = 'PASAR AL FOLIO:________'
#    para = document.add_paragraph("ESTA PÁGINA ESTÁ DEDICADA A: ")
#    para.add_run(i)

    #https://www.youtube.com/watch?v=W_Au8bLo82E
    #https://stackoverflow.com/questions/53194725/python-docx-table-row-height
    menuTable = document.add_table(22,12)
    menuTable.width = Cm(32.69)
    for cell in menuTable.columns[0].cells:
        cell.width = Cm(1.24)
    for cell in menuTable.columns[1].cells:
        cell.width = Cm(1.45)
    for cell in menuTable.columns[2].cells:
        cell.width = Cm(3.3)
    for cell in menuTable.columns[3].cells:
        cell.width = Cm(6.7)
    for cell in menuTable.columns[4].cells:
        cell.width = Cm(6.75)
    for cell in menuTable.columns[5].cells:
        cell.width = Cm(1.75)
    for cell in menuTable.columns[6].cells:
        cell.width = Cm(2)
    for cell in menuTable.columns[7].cells:
        cell.width = Cm(1.75)
    for cell in menuTable.columns[8].cells:
        cell.width = Cm(1.75)
    for cell in menuTable.columns[9].cells:
        cell.width = Cm(1.75)
    for cell in menuTable.columns[10].cells:
        cell.width = Cm(1.5)
    for cell in menuTable.columns[11].cells:
        cell.width = Cm(2.75)
    for row in menuTable.rows:
        row.height = Cm(0.7)
    menuTable.style = 'Table Grid'
    hdr_Cells = menuTable.rows[0].cells
    hdr_Cells[0].text = 'AÑO'
    hdr_Cells[1].text = 'FECHA'
    hdr_Cells[2].text = 'PROCEDENCIA'
    hdr_Cells[3].text = 'MEDICO'
    hdr_Cells[4].text = 'DIRECCIÓN'
    hdr_Cells[5].text = 'CED. PR.'
    hdr_Cells[6].text = 'No. FAC.'
    hdr_Cells[7].text = 'No. REC.'
    hdr_Cells[8].text = 'C. ADQ.'
    hdr_Cells[9].text = 'C. VEN.'
    hdr_Cells[10].text = 'SALDO'
    hdr_Cells[11].text = 'OBSER'
    menuTable.cell(1,0).text = '2020'
    menuTable.cell(1,1).text = 'ENE 1'
    menuTable.cell(1,10).text = str(lista_conjunta[cont][2])
    
    document.add_page_break()
    
    cont+=1

#ESPECIFICAR
document.save('Demo.docx')

